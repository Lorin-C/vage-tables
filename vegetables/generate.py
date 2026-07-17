from flask import Flask, request, send_file
from docx import Document
from docx.table import Table
from docx.shared import Pt
from docx.enum.text import WD_TAB_ALIGNMENT
from datetime import datetime
import os


app = Flask(__name__)


BASE_DIR = os.path.dirname(
    os.path.dirname(
        os.path.abspath(__file__)
    )
)


TEMPLATE = os.path.join(
    BASE_DIR,
    "uploads",
    "form2026.docx"
)

# =========================
# 字体设置
# =========================

def set_font(run, size, bold=True):

    run.font.name = "宋体"

    run._element.rPr.rFonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia",
        "宋体"
    )

    run.font.size = Pt(size)

    run.bold = bold

# =========================
# 获取表格
# =========================

def get_all_tables(doc):

    tables=[]


    for tbl in doc._body._body.iter(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl"
    ):

        tables.append(
            Table(tbl,doc._body)
        )


    return tables

# =========================
# Vercel API入口
# =========================

@app.route(
    "/api/generate",
    methods=["POST"]
)
def generate():


    customer=request.form.get(
        "customer",
        ""
    )


    # 最大6字

    customer=customer[:6]



    date=datetime.now().strftime(
        "%Y年%m月%d日"
    )



    text=request.form.get(
        "dishes",
        ""
    )


    dishes=[

        x.strip()

        for x in text.splitlines()

        if x.strip()

    ]



    doc=Document(
        TEMPLATE
    )



    tables=get_all_tables(
        doc
    )

    # =========================
    # 客户 日期
    # =========================

    if len(tables)>0:


        header=tables[0]


        for row in header.rows:


            for cell in row.cells:


                for p in cell.paragraphs:


                    if (
                        "客户：" in p.text
                    ):


                        for run in p.runs:

                            run.text=""



                        p.paragraph_format.tab_stops.add_tab_stop(
                            Pt(300),
                            WD_TAB_ALIGNMENT.RIGHT
                        )



                        run=p.add_run(

                            "客户："

                            +

                            customer

                            +

                            "\t"

                            +

                            date

                        )


                        set_font(
                            run,
                            14,
                            True
                        )


                        break

    # =========================
    # 菜品
    # =========================

    dish_cells=[]



    for table in tables[1:]:


        for row in table.rows:


            cells=row.cells


            if len(cells)>=5:


                number=cells[0].text.strip()



                if number.isdigit():


                    n=int(number)


                    if 1<=n<=47:


                        dish_cells.append(
                            cells[1]
                        )





    for i,cell in enumerate(dish_cells):


        if i<len(dishes):


            p=cell.paragraphs[0]


            for run in p.runs:

                run.text=""



            run=p.add_run(
                dishes[i]
            )


            set_font(
                run,
                18,
                True
            )

    # =========================
    # Vercel临时目录
    # =========================

    output="/tmp/蔬菜单.docx"



    doc.save(
        output
    )



    return send_file(

        output,

        as_attachment=True,

        download_name="蔬菜单.docx"

    )
